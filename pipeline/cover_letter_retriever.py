# Cover Letter Retriever - Script 5
# Checks if batch is complete, downloads cover letters, saves as .docx files
#
# Required environment variables:
#   ANTHROPIC_API_KEY  - your Anthropic API key
#   BASE_PATH          - folder where pending_batch.json is and output folders will be created
#
# Input: pending_batch.json in BASE_PATH (created by cover_letter_generator.py)
# Output: Cover_Letters/[date]/ folder with .docx files

import subprocess
import sys

for pkg in ["anthropic", "python-docx"]:
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "--quiet"])

import os
import csv
import json
import re
from datetime import datetime
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# CONFIGURATION
# ============================================================

BASE_PATH = os.environ.get("BASE_PATH", os.getcwd())
PENDING_BATCH_FILE = os.path.join(BASE_PATH, "pending_batch.json")
COVER_LETTER_HISTORY = os.path.join(BASE_PATH, "cover_letters_history.csv")
COVER_LETTERS_FOLDER = os.path.join(BASE_PATH, "Cover_Letters")

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def sanitize_filename(text):
    """Remove characters that are not safe for Windows filenames."""
    safe = re.sub(r'[<>:"/\\|?*]', '_', text)
    safe = safe.strip(' .')
    if len(safe) > 80:
        safe = safe[:80]
    return safe


def save_cover_letter_as_docx(cover_letter_text, filepath):
    """Save cover letter text as a formatted .docx file.
    Justified alignment. No extra spacing between paragraphs.
    Only a blank line before Best regards."""
    doc = Document()

    # Set default font and paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = cover_letter_text.strip().split('\n')

    for line in lines:
        stripped = line.strip()
        if stripped == '':
            p = doc.add_paragraph('')
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(filepath)


def update_cover_letter_history(jobs_with_letters):
    """Add jobs to the cover letter history file."""
    file_exists = os.path.exists(COVER_LETTER_HISTORY)

    fieldnames = ["job_url", "title", "company", "scope", "score", "model",
                   "docx_path", "date_generated"]

    with open(COVER_LETTER_HISTORY, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        if not file_exists:
            writer.writeheader()
        for job in jobs_with_letters:
            writer.writerow(job)


# ============================================================
# MAIN SCRIPT
# ============================================================

def main():
    print("=" * 60)
    print("COVER LETTER RETRIEVER")
    print("=" * 60)

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("\nERROR: ANTHROPIC_API_KEY not found!")
        print("Set it with: setx ANTHROPIC_API_KEY \"sk-ant-...\"")
        print("Then restart Command Prompt and try again.")
        sys.exit(1)

    if not os.path.exists(PENDING_BATCH_FILE):
        print(f"\nERROR: {PENDING_BATCH_FILE} not found!")
        print("Run the generator first: py -3.12 cover_letter_generator.py")
        sys.exit(1)

    with open(PENDING_BATCH_FILE, "r", encoding="utf-8") as f:
        batch_info = json.load(f)

    batch_id = batch_info["batch_id"]
    submitted_at = batch_info["submitted_at"]
    total_jobs = batch_info["total_jobs"]
    jobs_lookup = {j["custom_id"]: j for j in batch_info["jobs"]}

    print(f"\n  Batch ID: {batch_id}")
    print(f"  Submitted at: {submitted_at}")
    print(f"  Total jobs: {total_jobs}")
    print(f"  Sonnet: {batch_info['sonnet_count']}, Haiku: {batch_info['haiku_count']}")

    print(f"\nChecking batch status...")
    client = Anthropic()

    try:
        batch = client.messages.batches.retrieve(batch_id)
    except Exception as e:
        print(f"\nERROR checking batch: {e}")
        sys.exit(1)

    status = batch.processing_status
    print(f"  Status: {status}")

    if status == "in_progress":
        counts = batch.request_counts
        if counts:
            processing = counts.processing
            succeeded = counts.succeeded
            errored = counts.errored
            total = processing + succeeded + errored
            print(f"  Progress: {succeeded} succeeded, {processing} processing, {errored} errored (of {total})")
        print(f"\nBatch is still processing. Try again later.")
        print(f"Run this script again: py -3.12 cover_letter_retriever.py")
        sys.exit(0)

    elif status == "ended":
        counts = batch.request_counts
        succeeded = counts.succeeded if counts else 0
        errored = counts.errored if counts else 0
        print(f"  Batch complete! {succeeded} succeeded, {errored} errored")

    else:
        print(f"\nUnexpected batch status: {status}")
        print("Check your Anthropic dashboard: https://console.anthropic.com")
        sys.exit(1)

    # Create output folder
    today = datetime.now().strftime("%Y-%m-%d")
    output_folder = os.path.join(COVER_LETTERS_FOLDER, today)
    os.makedirs(output_folder, exist_ok=True)
    print(f"\n  Saving cover letters to: {output_folder}")

    print(f"\nDownloading results...")

    success_count = 0
    error_count = 0
    jobs_with_letters = []

    try:
        for result in client.messages.batches.results(batch_id):
            custom_id = result.custom_id
            job_info = jobs_lookup.get(custom_id, {})

            title = job_info.get("title", "Unknown")
            company = job_info.get("company", "Unknown")
            score = job_info.get("score", 0)
            scope_name = job_info.get("scope_name", "Unknown")
            job_url = job_info.get("url", "")
            model = job_info.get("model", "unknown")
            model_tag = job_info.get("model_tag", "H")

            if result.result.type == "succeeded":
                message = result.result.message
                cover_letter_text = ""
                for block in message.content:
                    if block.type == "text":
                        cover_letter_text += block.text

                if not cover_letter_text.strip():
                    print(f"  WARNING: Empty response for {title} @ {company}")
                    error_count += 1
                    continue

                # Score-based filename: 85S-Title - Company.docx
                safe_title = sanitize_filename(title)
                safe_company = sanitize_filename(company)
                filename = f"{score}{model_tag}-{safe_title} - {safe_company}.docx"
                filepath = os.path.join(output_folder, filename)

                # Handle duplicate filenames
                counter = 1
                while os.path.exists(filepath):
                    filename = f"{score}{model_tag}-{safe_title} - {safe_company}_{counter}.docx"
                    filepath = os.path.join(output_folder, filename)
                    counter += 1

                save_cover_letter_as_docx(cover_letter_text, filepath)

                success_count += 1
                print(f"  [{success_count}] {score}{model_tag} - {title} @ {company} ({scope_name})")

                jobs_with_letters.append({
                    "job_url": job_url,
                    "title": title,
                    "company": company,
                    "scope": scope_name,
                    "score": score,
                    "model": model,
                    "docx_path": filepath,
                    "date_generated": today
                })

            else:
                error_type = result.result.type
                print(f"  ERROR for {title} @ {company}: {error_type}")
                error_count += 1

    except Exception as e:
        print(f"\nERROR downloading results: {e}")
        print("Some letters may have been saved. Check the folder.")

    if jobs_with_letters:
        print(f"\nUpdating cover letter history...")
        update_cover_letter_history(jobs_with_letters)
        print(f"  Added {len(jobs_with_letters)} entries to {COVER_LETTER_HISTORY}")

    print(f"\n{'=' * 60}")
    print("SUMMARY")
    print(f"{'=' * 60}")
    print(f"  Cover letters saved: {success_count}")
    print(f"  Errors: {error_count}")
    print(f"  Folder: {output_folder}")
    print(f"  Estimated cost: ${batch_info['estimated_cost']:.4f}")

    if success_count > 0:
        print(f"\n  Your cover letters are ready!")
        print(f"  Open the folder: {output_folder}")

        done_file = os.path.join(BASE_PATH, f"batch_done_{today}.json")
        try:
            os.rename(PENDING_BATCH_FILE, done_file)
            print(f"\n  Batch file moved to: {done_file}")
        except Exception:
            pass

    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
